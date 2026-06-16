"""生成 annual_financials_v2 字段来源说明 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- 页面设置 --
for section in doc.sections:
    section.orientation = 1  # 横向
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)

# -- 标题 --
title = doc.add_heading('annual_financials_v2 字段来源说明', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('数据表：Supabase → annual_financials_v2')
doc.add_paragraph('数据时间范围：2016-2025 年')
doc.add_paragraph('')

# ============================================================
# 辅助函数
# ============================================================
def set_cell_font(cell, text, bold=False, size=9, color=None):
    """设置单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_table_with_header(doc, title_text, headers, rows):
    """添加带标题和表头的表格"""
    doc.add_heading(title_text, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True, size=9)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            set_cell_font(table.rows[r + 1].cells[c], val, size=9)

    doc.add_paragraph('')
    return table

# ============================================================
# 一、东方财富 Datacenter API
# ============================================================
doc.add_heading('一、东方财富 Datacenter API 直接获取', level=1)
doc.add_paragraph('接口：https://datacenter.eastmoney.com/securities/api/data/v1/get')
doc.add_paragraph('共 15 个字段，来自三个报表：利润表(RPT_DMSK_FN_INCOME)、资产负债表(RPT_DMSK_FN_BALANCE)、现金流量表(RPT_DMSK_FN_CASHFLOW)')

dc_headers = ['#', '字段名', '中文含义', '来源报表', '原始列名']
dc_rows = [
    ['1', 'operating_revenue', '营业收入', 'RPT_DMSK_FN_INCOME', 'TOTAL_OPERATE_INCOME'],
    ['2', 'operating_cost', '营业成本', 'RPT_DMSK_FN_INCOME', 'TOTAL_OPERATE_COST'],
    ['3', 'selling_expense', '销售费用', 'RPT_DMSK_FN_INCOME', 'SALE_EXPENSE'],
    ['4', 'admin_expense', '管理费用', 'RPT_DMSK_FN_INCOME', 'MANAGE_EXPENSE'],
    ['5', 'finance_expense', '财务费用', 'RPT_DMSK_FN_INCOME', 'FINANCE_EXPENSE'],
    ['6', 'operating_profit', '营业利润', 'RPT_DMSK_FN_INCOME', 'OPERATE_PROFIT'],
    ['7', 'total_profit', '利润总额', 'RPT_DMSK_FN_INCOME', 'TOTAL_PROFIT'],
    ['8', 'parent_net_profit', '归母净利润', 'RPT_DMSK_FN_INCOME', 'PARENT_NETPROFIT'],
    ['9', 'deducted_net_profit', '扣非净利润', 'RPT_DMSK_FN_INCOME', 'DEDUCT_PARENT_NETPROFIT'],
    ['10', 'inventory', '存货', 'RPT_DMSK_FN_BALANCE', 'INVENTORY'],
    ['11', 'fixed_asset', '固定资产', 'RPT_DMSK_FN_BALANCE', 'FIXED_ASSET'],
    ['12', 'cash_and_equivalents', '现金及现金等价物（货币资金）', 'RPT_DMSK_FN_BALANCE', 'MONETARYFUNDS'],
    ['13', 'advance_contract_liab', '预收款/合同负债（初次写入）', 'RPT_DMSK_FN_BALANCE', 'ADVANCE_RECEIVABLES'],
    ['14', 'operating_cash_flow', '经营活动现金流净额', 'RPT_DMSK_FN_CASHFLOW', 'NETCASH_OPERATE'],
    ['15', 'capex', '资本开支（购建固定资产等）', 'RPT_DMSK_FN_CASHFLOW', 'CONSTRUCT_LONG_ASSET'],
]
add_table_with_header(doc, '1.1 东方财富 Datacenter API 字段（15个）', dc_headers, dc_rows)

# ============================================================
# 二、新浪财经 API
# ============================================================
doc.add_heading('二、新浪财经 API 直接获取', level=1)
doc.add_paragraph('接口：https://quotes.sina.cn/cn/api/openapi.php/CompanyFinanceService.getFinanceReport2022')
doc.add_paragraph('共 12 个字段，来自资产负债表和利润表')

sina_headers = ['#', '字段名', '中文含义', '来源报表', 'Sina 字段名']
sina_rows = [
    ['1', 'construction_in_progress', '在建工程', '资产负债表 (fzb)', '在建工程'],
    ['2', 'goodwill', '商誉', '资产负债表 (fzb)', '商誉'],
    ['3', 'short_term_loan', '短期借款', '资产负债表 (fzb)', '短期借款'],
    ['4', 'long_term_loan', '长期借款', '资产负债表 (fzb)', '长期借款'],
    ['5', 'bonds_payable', '应付债券', '资产负债表 (fzb)', '应付债券'],
    ['6', 'total_shares', '总股本（实收资本）', '资产负债表 (fzb)', '实收资本(或股本)'],
    ['7', 'total_assets', '总资产', '资产负债表 (fzb)', '资产总计'],
    ['8', 'total_liabilities', '总负债', '资产负债表 (fzb)', '负债合计'],
    ['9', 'parent_equity', '归母股东权益（净资产）', '资产负债表 (fzb)', '归属于母公司股东权益合计'],
    ['10', 'advance_contract_liab', '预收款/合同负债（覆盖写入）', '资产负债表 (fzb)', '2017前=预收款项, 2017后=合同负债'],
    ['11', 'rd_expense', '研发费用', '利润表 (lrb)', '研发费用'],
]
add_table_with_header(doc, '2.1 新浪财经 API 字段（12个）', sina_headers, sina_rows)

# ============================================================
# 三、旧表迁移
# ============================================================
doc.add_heading('三、旧表 (annual_financials) 迁移字段', level=1)
doc.add_paragraph('来源：v1 阶段累积的 annual_financials 表，原始 API 来源已不可追溯。共 9 个字段。')

migrate_headers = ['#', '字段名', '中文含义', '旧表字段名', '备注']
migrate_rows = [
    ['1', 'total_revenue', '营业总收入', 'total_revenue', ''],
    ['2', 'gross_profit', '毛利润', 'gross_profit', ''],
    ['3', 'gross_margin', '毛利率 (%)', 'gross_margin', ''],
    ['4', 'net_margin', '净利率 (%)', 'net_margin', ''],
    ['5', 'roe', '净资产收益率 (%)', 'roe', ''],
    ['6', 'roa', '总资产报酬率 (%)', 'roa', ''],
    ['7', 'total_market_cap', '总市值', 'market_cap', '大部分为 NULL'],
    ['8', 'revenue_yoy', '营业总收入同比增长率 (%)', 'revenue_growth', ''],
    ['9', 'net_profit_yoy', '归母净利润同比增长率 (%)', 'net_profit_growth', ''],
]
add_table_with_header(doc, '3.1 旧表迁移字段（9个）', migrate_headers, migrate_rows)

# ============================================================
# 四、计算字段
# ============================================================
doc.add_heading('四、计算字段（派生字段）', level=1)
doc.add_paragraph('来源：Python 脚本 calc_derived_fields.py，从已有字段计算得出。共 15 个字段。')
doc.add_paragraph('计算规则：NULL 源字段视为 0；除法分母为 0 或 NULL 时结果设为 NULL；parent_net_profit 为 NULL 时有效资产收益率、净利润率、每股净利润均置 NULL；gross_profit 为 0 时全部费用占毛利比置 NULL。')

calc_headers = ['#', '字段名', '中文含义', '计算公式', '依赖字段']
calc_rows = [
    ['1', 'effective_asset_return', '有效资产收益率 (%)',
     'parent_net_profit / (total_assets − cash_and_equivalents − goodwill) × 100',
     'parent_net_profit, total_assets, cash_and_equivalents, goodwill'],
    ['2', 'hard_profit', '硬朗度利润（自由现金流）',
     'operating_cash_flow − capex',
     'operating_cash_flow, capex'],
    ['3', 'anchor_assets', '锚定资产',
     'fixed_asset + construction_in_progress + inventory',
     'fixed_asset, construction_in_progress, inventory'],
    ['4', 'anchor_asset_ratio', '锚定资产占比 (%)',
     '(fixed_asset + construction_in_progress + inventory) / total_assets × 100',
     'fixed_asset, construction_in_progress, inventory, total_assets'],
    ['5', 'net_profit_margin', '净利润率 (%)',
     'parent_net_profit / operating_revenue × 100',
     'parent_net_profit, operating_revenue'],
    ['6', 'eps_calculated', '每股净利润',
     'parent_net_profit / total_shares',
     'parent_net_profit, total_shares'],
    ['7', 'selling_to_gross', '销售费用占毛利比 (%)',
     'selling_expense / gross_profit × 100',
     'selling_expense, gross_profit'],
    ['8', 'admin_to_gross', '管理费用占毛利比 (%)',
     'admin_expense / gross_profit × 100',
     'admin_expense, gross_profit'],
    ['9', 'selling_admin_to_gross', '销售+管理费用占毛利比 (%)',
     '(selling_expense + admin_expense) / gross_profit × 100',
     'selling_expense, admin_expense, gross_profit'],
    ['10', 'rd_to_gross', '研发费用占毛利比 (%)',
     'rd_expense / gross_profit × 100',
     'rd_expense, gross_profit'],
    ['11', 'sga_rd_to_gross', '销管研三费占毛利比 (%)',
     '(selling_expense + admin_expense + rd_expense) / gross_profit × 100',
     'selling_expense, admin_expense, rd_expense, gross_profit'],
    ['12', 'finance_to_gross', '财务费用占毛利比 (%)',
     'finance_expense / gross_profit × 100',
     'finance_expense, gross_profit'],
    ['13', 'advance_to_revenue', '预收款占总营收比 (%)',
     'advance_contract_liab / operating_revenue × 100',
     'advance_contract_liab, operating_revenue'],
    ['14', 'debt_ratio_ex_advance', '剔除预收款后的资产负债率 (%)',
     '(total_liabilities − advance_contract_liab) / (total_assets − advance_contract_liab) × 100',
     'total_liabilities, advance_contract_liab, total_assets'],
    ['15', 'debt_equity_ex_cash', '剔除预收款后的债务股权比',
     '(short_term_loan + long_term_loan + bonds_payable − cash_and_equivalents) / parent_equity',
     'short_term_loan, long_term_loan, bonds_payable, cash_and_equivalents, parent_equity'],
]
add_table_with_header(doc, '4.1 计算字段（15个）', calc_headers, calc_rows)

# ============================================================
# 五、汇总
# ============================================================
doc.add_heading('五、汇总', level=1)

summary_headers = ['数据来源', '字段数', '字段列表']
summary_rows = [
    ['东方财富 Datacenter API', '15',
     'operating_revenue, operating_cost, selling_expense, admin_expense, finance_expense, operating_profit, total_profit, parent_net_profit, deducted_net_profit, inventory, fixed_asset, cash_and_equivalents, advance_contract_liab(初次), operating_cash_flow, capex'],
    ['新浪财经 API', '12',
     'construction_in_progress, goodwill, short_term_loan, long_term_loan, bonds_payable, total_shares, total_assets, total_liabilities, parent_equity, advance_contract_liab(覆盖), rd_expense'],
    ['旧表迁移', '9',
     'total_revenue, gross_profit, gross_margin, net_margin, roe, roa, total_market_cap, revenue_yoy, net_profit_yoy'],
    ['计算字段', '15',
     'effective_asset_return, hard_profit, anchor_assets, anchor_asset_ratio, net_profit_margin, eps_calculated, selling_to_gross, admin_to_gross, selling_admin_to_gross, rd_to_gross, sga_rd_to_gross, finance_to_gross, advance_to_revenue, debt_ratio_ex_advance, debt_equity_ex_cash'],
    ['系统字段', '7',
     'id, code, stock_name, year, data_source, created_at, updated_at'],
    ['总计', '58', '（含 3 个未填充字段：eps, bps, 以及部分为空的 total_market_cap）'],
]
add_table_with_header(doc, '5.1 字段来源汇总', summary_headers, summary_rows)

# ============================================================
# 六、当前为空字段
# ============================================================
doc.add_heading('六、当前未填充字段', level=1)
empty_headers = ['字段名', '中文含义', '状态']
empty_rows = [
    ['eps', '基本每股收益', '全部为 NULL，暂无数据源'],
    ['bps', '每股净资产', '全部为 NULL，暂无数据源'],
    ['total_market_cap', '总市值', '仅旧表迁移时有值，大部分为 NULL'],
]
add_table_with_header(doc, '6.1 未填充字段', empty_headers, empty_rows)

# -- 保存 --
output_path = os.path.join(os.path.dirname(__file__), '..', 'annual_financials_v2_字段来源说明.docx')
doc.save(output_path)
print(f'✅ 文档已保存到: {os.path.abspath(output_path)}')